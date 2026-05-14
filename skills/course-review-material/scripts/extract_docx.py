"""Extract content from .docx files to structured Markdown.

Extracts text, embedded images, and tables.
Preserves heading styles, bold/italic formatting, and lists.

Usage:
    python extract_docx.py <input.docx> [--output-dir <dir>] [--extract-images]
"""

import argparse
import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

def extract_images_from_docx(doc, output_dir):
    images_dir = os.path.join(output_dir, "images")
    os.makedirs(images_dir, exist_ok=True)
    rId_to_filename = {}
    img_counter = 0

    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image = rel.target_part
            ext = os.path.splitext(image.partname)[1]
            if not ext:
                ext = ".png"
            img_counter += 1
            img_name = f"img_{img_counter:03d}{ext}"
            img_path = os.path.join(images_dir, img_name)
            with open(img_path, "wb") as f:
                f.write(image.blob)
            rId_to_filename[rel.rId] = img_name

    return rId_to_filename


def get_heading_level(paragraph):
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name.startswith("Heading") or style_name.startswith("heading"):
        try:
            level = int(style_name.split()[-1])
            return min(level, 6)
        except (ValueError, IndexError):
            return 1
    if style_name == "Title":
        return 1
    return 0


def format_run(run):
    text = run.text
    if run.bold and run.italic:
        return f"***{text}***"
    if run.bold:
        return f"**{text}**"
    if run.italic:
        return f"*{text}*"
    return text


def is_list_paragraph(paragraph):
    style_name = paragraph.style.name.lower() if paragraph.style else ""
    if "list" in style_name:
        return True
    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is not None:
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            return True
    return False


def convert_table(table):
    rows = []
    for row in table.rows:
        cells = [cell.text.replace("\n", " ") for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    md_lines = []
    md_lines.append("| " + " | ".join(rows[0]) + " |")
    md_lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
    for row in rows[1:]:
        padded = row + [""] * (len(rows[0]) - len(row))
        md_lines.append("| " + " | ".join(padded[: len(rows[0])]) + " |")

    return "\n".join(md_lines)


def _get_para_image_refs(paragraph, rId_to_filename):
    refs = []
    drawings = paragraph._element.findall(".//" + qn("w:drawing"))
    for drawing in drawings:
        blips = drawing.findall(".//" + qn("a:blip"))
        for blip in blips:
            rId = blip.get(qn("r:embed"))
            if rId and rId in rId_to_filename:
                refs.append(rId_to_filename[rId])
    return refs


def extract_docx(filepath, output_dir, extract_images=True):
    doc = Document(filepath)

    os.makedirs(output_dir, exist_ok=True)
    text_dir = os.path.join(output_dir, "text")
    os.makedirs(text_dir, exist_ok=True)

    rId_to_filename = {}
    if extract_images:
        print("[extract] exporting embedded images...")
        rId_to_filename = extract_images_from_docx(doc, output_dir)

    raw_lines = []
    md_lines = []

    for para in doc.paragraphs:
        text = para.text
        raw_lines.append(text)

        heading_level = get_heading_level(para)
        if heading_level:
            prefix = "#" * heading_level
            md_lines.append(f"\n{prefix} {text}\n")
            image_refs = _get_para_image_refs(para, rId_to_filename)
            for img_name in image_refs:
                md_lines.append(f"![{img_name}](images/{img_name})\n")
                md_lines.append(f"<!-- image source: {img_name} -->\n")
            continue

        if is_list_paragraph(para):
            formatted_runs = "".join(format_run(r) for r in para.runs)
            md_lines.append(f"- {formatted_runs}")
            image_refs = _get_para_image_refs(para, rId_to_filename)
            for img_name in image_refs:
                md_lines.append(f"![{img_name}](images/{img_name})\n")
                md_lines.append(f"<!-- image source: {img_name} -->\n")
            continue

        if text.strip():
            formatted_runs = "".join(format_run(r) for r in para.runs)
            md_lines.append(f"\n{formatted_runs}\n")

        image_refs = _get_para_image_refs(para, rId_to_filename)
        for img_name in image_refs:
            md_lines.append(f"![{img_name}](images/{img_name})\n")
            md_lines.append(f"<!-- image source: {img_name} -->\n")

    for table in doc.tables:
        md_lines.append("\n" + convert_table(table) + "\n")

    raw_text_path = os.path.join(text_dir, "raw_text.txt")
    with open(raw_text_path, "w", encoding="utf-8") as f:
        f.write("\n".join(raw_lines))

    extracted_path = os.path.join(output_dir, "extracted.md")
    with open(extracted_path, "w", encoding="utf-8") as f:
        f.write("\n".join(md_lines))

    print(f"[done] raw text: {raw_text_path}")
    print(f"[done] extracted: {extracted_path}")
    if rId_to_filename:
        print(f"[done] {len(rId_to_filename)} images exported to {os.path.join(output_dir, 'images')}")
    return extracted_path


def main():
    parser = argparse.ArgumentParser(description="Extract docx content to structured Markdown")
    parser.add_argument("input", help="Input .docx file")
    parser.add_argument("--output-dir", "-o", default=None, help="Output directory (default: {input_name}/)")
    parser.add_argument("--no-images", action="store_true", help="Skip image extraction")
    args = parser.parse_args()

    input_path = Path(args.input).resolve()
    if not input_path.exists():
        print(f"Error: file not found: {input_path}")
        sys.exit(1)

    output_dir = args.output_dir or str(input_path.parent / input_path.stem)

    try:
        extract_docx(str(input_path), output_dir, extract_images=not args.no_images)
    except Exception as e:
        print(f"Error: failed to process {input_path}: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
